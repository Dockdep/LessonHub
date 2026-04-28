"""Smoke tests for binary-format extractors. We don't ship sample PDFs/EPUBs
in the repo — instead we either build a tiny synthetic DOCX in-memory (the
only format whose lib makes that easy) or trust the upstream lib for the
others."""
import io

import pytest

from tools import rag_extractors
from tools.rag_extractors import extract_docx, extract_pdf


def test_extract_docx_preserves_heading_levels_as_markdown():
    """Build a tiny DOCX in memory and verify Heading 1/2 become #/##."""
    import docx  # type: ignore[import-not-found]

    document = docx.Document()
    document.add_heading("Top Title", level=1)
    document.add_paragraph("Some intro text.")
    document.add_heading("A Subsection", level=2)
    document.add_paragraph("Body of the subsection.")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    text = extract_docx(buf.getvalue())

    # Heading 1 → "# Top Title", Heading 2 → "## A Subsection"
    assert "# Top Title" in text
    assert "## A Subsection" in text
    # Body paragraphs are present
    assert "Some intro text." in text
    assert "Body of the subsection." in text


def test_extract_docx_skips_empty_paragraphs():
    import docx  # type: ignore[import-not-found]

    document = docx.Document()
    document.add_paragraph("First.")
    document.add_paragraph("")  # explicitly empty
    document.add_paragraph("Third.")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    text = extract_docx(buf.getvalue())
    # No double blank-line gaps from empty paragraphs
    assert "First." in text
    assert "Third." in text
    # Two non-empty lines, joined by \n\n — exactly one blank line between them
    assert "First.\n\nThird." in text


def test_pypdf_wrapper_swallows_unicode_decode_error_at_init():
    """The pypdf wrapper must catch UnicodeDecodeError raised during
    PdfReader init (3-Heights Security Shell PDFs do this — non-UTF-8
    metadata) and return '' so the outer fallback can kick in. Without this
    catch, the FastAPI ValueError handler would turn it into a confusing
    400 BadRequest with the raw codec message."""
    # Random bytes that aren't a valid PDF — pypdf will raise during init.
    garbage = b"\x00\xe2\x80\x93 not a real pdf header at all"
    # Must return "" rather than propagate.
    assert rag_extractors._extract_pdf_pypdf(garbage) == ""


def test_extract_pdf_falls_back_to_pdfminer_when_pypdf_yields_nothing(monkeypatch):
    """Outer extract_pdf glues pypdf and pdfminer together: when pypdf
    returns empty (covers both 'opened but extracted no text' and 'failed
    to open'), pdfminer's output is used."""
    monkeypatch.setattr(rag_extractors, "_extract_pdf_pypdf", lambda data: "")
    monkeypatch.setattr(rag_extractors, "_extract_pdf_pdfminer",
                        lambda data: "fallback text from pdfminer")

    assert extract_pdf(b"%PDF-1.5\n...") == "fallback text from pdfminer"


def test_extract_pdf_raises_clear_error_when_both_extractors_fail(monkeypatch):
    """Both extractors returning empty should produce a human-readable
    ValueError, not the raw codec error from whichever lib failed last."""
    monkeypatch.setattr(rag_extractors, "_extract_pdf_pypdf", lambda data: "")
    monkeypatch.setattr(rag_extractors, "_extract_pdf_pdfminer", lambda data: "")

    with pytest.raises(ValueError, match="Could not extract any text"):
        extract_pdf(b"%PDF-1.5\nbroken")


def test_extract_pdf_uses_pypdf_when_it_succeeds(monkeypatch):
    """Happy path: pdfminer fallback shouldn't run if pypdf already got text."""
    pdfminer_calls = []
    monkeypatch.setattr(rag_extractors, "_extract_pdf_pypdf", lambda data: "from pypdf")
    monkeypatch.setattr(
        rag_extractors,
        "_extract_pdf_pdfminer",
        lambda data: pdfminer_calls.append(1) or "from pdfminer",
    )

    assert extract_pdf(b"%PDF-1.5\n") == "from pypdf"
    assert pdfminer_calls == []  # fallback never ran
